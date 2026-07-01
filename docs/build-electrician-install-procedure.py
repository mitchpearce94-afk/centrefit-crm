"""Generate the electrician install procedure as an editable .docx.

Photos are represented as bordered placeholder boxes captioned with the
shot Mitchell wants there — open in Word, click the box, Insert > Picture
to drop the real photo on top. Each placeholder is in its own table cell
so it's easy to resize without breaking the surrounding flow.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Styling helpers ────────────────────────────────────────────────────────

CENTREFIT_BLUE = RGBColor(0x3B, 0x82, 0xF6)
WARNING_RED = RGBColor(0xC0, 0x39, 0x2B)
MUTED = RGBColor(0x64, 0x74, 0x8B)


def shade_cell(cell, hex_colour: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


def set_cell_border(cell, hex_colour: str = "94A3B8", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), hex_colour)
        borders.append(el)
    tc_pr.append(borders)


def photo_placeholder(doc: Document, description: str, height_cm: float = 6.5) -> None:
    """Bordered single-cell table that says '[Insert photo: description]'.
    Right-click → Insert → Picture inside the cell to drop the real shot."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade_cell(cell, "F1F5F9")
    set_cell_border(cell, "94A3B8")
    # Force row height
    tr_pr = cell._tc.getparent().get_or_add_trPr()
    trh = OxmlElement("w:trHeight")
    trh.set(qn("w:val"), str(int(height_cm * 567)))  # twips per cm
    trh.set(qn("w:hRule"), "atLeast")
    tr_pr.append(trh)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    icon = p.add_run("📷  ")
    icon.font.size = Pt(14)
    label = p.add_run(f"Insert photo: {description}")
    label.italic = True
    label.font.size = Pt(11)
    label.font.color.rgb = MUTED

    doc.add_paragraph()  # spacer


def heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = CENTREFIT_BLUE if level == 1 else RGBColor(0x0F, 0x17, 0x2A)


def numbered(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def warning(doc: Document, text: str) -> None:
    """Yellow-shaded full-width warning callout."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    shade_cell(cell, "FEF3C7")
    set_cell_border(cell, "F59E0B", size="12")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    icon = p.add_run("⚠  ")
    icon.bold = True
    icon.font.color.rgb = WARNING_RED
    icon.font.size = Pt(12)
    body = p.add_run(text)
    body.bold = True
    body.font.size = Pt(10.5)
    body.font.color.rgb = RGBColor(0x78, 0x35, 0x0F)
    doc.add_paragraph()


def intro_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = MUTED


# ── Build the doc ──────────────────────────────────────────────────────────

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Page margins (a bit tighter than default so screenshots/photos sit larger)
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)


# ── Title page ─────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("Centrefit Group")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = MUTED

t = doc.add_paragraph()
t_run = t.add_run("Electrician Install Procedure")
t_run.bold = True
t_run.font.size = Pt(28)
t_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

sub = doc.add_paragraph()
sub_run = sub.add_run("Alarm Panel · Server Rack · Data · CCTV · Automation · Router · Nightlife")
sub_run.font.size = Pt(12)
sub_run.font.color.rgb = MUTED

intro_para(
    doc,
    "Step-by-step procedure for the electrician on a Centrefit fit-out. Work through the "
    "sections in order. Photo placeholders show what a finished install should look like — "
    "if your install doesn't match the photo, stop and check with the project manager before "
    "moving on.",
)

doc.add_paragraph()


# ── 1. ALARM PANEL ─────────────────────────────────────────────────────────

heading(doc, "1. Alarm Panel", level=1)

heading(doc, "Mount and fit the alarm panel", level=2)
numbered(doc, "Mount the alarm panel in the location shown on the plans.")
numbered(doc, "Wire each cable into its correct labelled terminal — do not improvise terminations.")
numbered(doc, "Confirm every cable is labelled correctly at both ends before powering on.")
photo_placeholder(doc, "Mounted alarm panel with all cables labelled and terminated")

heading(doc, "Coax termination in the roof space", level=2)
numbered(doc, "Terminate and split the coax in the roof space.")
numbered(doc, "Run cables from the splitters down to the integration wall plate, alongside the antenna cable.")
warning(doc, "ONE feed cable per ONE splitter. Do not daisy-chain feeds across splitters.")
photo_placeholder(doc, "Roof-space splitters with feed cables labelled and running to integration wall plate")


# ── 2. SERVER RACK ────────────────────────────────────────────────────────

heading(doc, "2. Server Rack", level=1)

numbered(doc, "Pull the rack off the wall by approximately 1.5 metres — enough clearance to service behind it.")
numbered(doc, "Mount the faceplates to the wall in their labelled positions.")
numbered(doc, "If the site has more than 16 cameras, move the additional cameras (cam 17 onwards) onto the data wall plate.")
numbered(
    doc,
    "Move the alarm panel data integration onto the data wall plate. See page 5 of the plans "
    "for the integration cable run.",
)
numbered(doc, "Run all data cables into the patch panel.")
photo_placeholder(doc, "Server rack pulled off wall, faceplates mounted, data cabling tidied to patch panel")


# ── 3. DATA ───────────────────────────────────────────────────────────────

heading(doc, "3. Data — Patch Panel Layout", level=1)

intro_para(doc, "Patch the following devices to the listed ports. Mis-patching here breaks the whole site.")

# Patch panel table
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr_cells = table.rows[0].cells
for i, h in enumerate(("Patch Panel", "Port", "Device")):
    p = hdr_cells[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    shade_cell(hdr_cells[i], "E2E8F0")

patch_rows = [
    ("Patch 1", "Port 1", "NBN"),
    ("Patch 1", "Port 2", "Alarm panel"),
    ("Patch 1", "Port 3", "NVR"),
    ("Patch 1", "Port 4", "Felixgate camera"),
    ("Patch 1", "Port 5", "Felixgate sensor"),
    ("Patch 1", "Port 6 onwards", "WAPs (Wi-Fi access points) — one per port"),
    ("Patch 2", "Port 1 onwards", "Cameras 17+ (overflow when site has > 16 cameras)"),
]
for panel, port, device in patch_rows:
    row = table.add_row().cells
    row[0].text = panel
    row[1].text = port
    row[2].text = device

doc.add_paragraph()
photo_placeholder(doc, "Patch panel fully terminated with port labels visible")


# ── 4. CCTV ───────────────────────────────────────────────────────────────

heading(doc, "4. CCTV", level=1)

numbered(doc, "Wire cameras 1 through 16 directly into the NVR.")
numbered(doc, "Confirm each camera channel matches the labelled camera location.")
photo_placeholder(doc, "NVR rear with cameras 1–16 connected and labelled")


# ── 5. AUTOMATION ─────────────────────────────────────────────────────────

heading(doc, "5. Automation", level=1)

heading(doc, "Antenna & coax", level=2)
numbered(doc, "Wire the antenna into the splitter at the labelled port.")
numbered(doc, "Wire the feed cables into the 8-port active tap.")
photo_placeholder(doc, "Splitter and 8-port active tap with antenna + feeds terminated and labelled")

heading(doc, "Power supplies", level=2)
warning(
    doc,
    "Polarity matters. If the alarm panel or access control power supply is wired the wrong "
    "way around, you WILL blow up the controller. Double-check labels before powering on.",
)
numbered(doc, "Wire in the alarm panel power supply — verify the labelled polarity is correct.")
numbered(doc, "Wire in the access control power supply — verify the labelled polarity is correct.")
numbered(
    doc,
    "Plug both power supplies into the powerboard that is plugged into the UPS at the bottom "
    "of the rack — not directly into the wall.",
)
photo_placeholder(doc, "Powerboard inside the UPS with alarm panel + access control supplies plugged in")

heading(doc, "Audio", level=2)
numbered(doc, "Wire the speaker cable into the amplifier.")
numbered(
    doc,
    "If a second 120W amplifier is fitted, wire the second zone into its 100V speaker output.",
)
numbered(
    doc,
    "If a WiiM player is fitted, unbox it and wire the supplied RCA plugs into Line 3 of the "
    "120W amplifier — the cables will be labelled.",
    bold_prefix="WiiM (if fitted): ",
)
photo_placeholder(doc, "Amplifier with speaker zones + WiiM RCA inputs wired and labelled")


# ── 6. ROUTER ─────────────────────────────────────────────────────────────

heading(doc, "6. Router", level=1)
numbered(doc, "Install the UniFi router on the top shelf of the rack.")
photo_placeholder(doc, "UniFi router mounted on the top shelf of the rack")


# ── 7. NIGHTLIFE ──────────────────────────────────────────────────────────

heading(doc, "7. Nightlife", level=1)
numbered(
    doc,
    "Install the powerboard from the Nightlife kit and plug it into the battery / power-surge "
    "GPO of the UPS.",
)
numbered(doc, "Install the Nightlife router on the top shelf of the rack.")
numbered(
    doc,
    "Run the Nightlife kiosk from the data wall plate into the port labelled 'Nightlife Only' "
    "on the router.",
)
numbered(
    doc,
    "Plug the client network into the patch panel, then through to the switch.",
)
numbered(doc, "Plug the Nightlife server into its dedicated port on the Nightlife router.")
photo_placeholder(doc, "Nightlife router top shelf with kiosk, client network, and server cabled")


# ── 8. TEST COMMISSION ───────────────────────────────────────────────────

heading(doc, "8. Test Commission", level=1)
intro_para(
    doc,
    "Final walk-through with the project manager. Confirm every system is up before leaving site.",
)
bullet(doc, "Alarm panel: power LED green, panel reports armed/disarmed correctly.")
bullet(doc, "Server rack: all data link lights up, no orphaned cables.")
bullet(doc, "CCTV: cameras 1–16 visible on NVR, cam 17+ (if any) reachable from the network.")
bullet(doc, "Automation: amplifier zones audible, WiiM (if fitted) streaming to Line 3.")
bullet(doc, "Router: UniFi reachable on management VLAN.")
bullet(doc, "Nightlife: kiosk live, client network reachable, Nightlife server up.")
bullet(doc, "All labels legible and final photos taken for the site Key Info tab.")
photo_placeholder(doc, "Final rack photo — door closed, labels visible, everything powered up")


# ── Sign-off ─────────────────────────────────────────────────────────────

doc.add_page_break()
heading(doc, "Sign-off", level=1)

sign_table = doc.add_table(rows=4, cols=2)
sign_table.style = "Light Grid Accent 1"

for i, label in enumerate(("Site", "Electrician", "Centrefit project manager", "Date")):
    sign_table.rows[i].cells[0].text = label
    p = sign_table.rows[i].cells[0].paragraphs[0]
    p.runs[0].bold = True
    shade_cell(sign_table.rows[i].cells[0], "F1F5F9")
    sign_table.rows[i].cells[1].text = ""

doc.add_paragraph()
notes_p = doc.add_paragraph()
notes_run = notes_p.add_run("Notes / Variances:")
notes_run.bold = True

for _ in range(6):
    p = doc.add_paragraph("_" * 110)
    p.runs[0].font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)


# ── Save ──────────────────────────────────────────────────────────────────

out_path = Path(__file__).parent / "Centrefit-Electrician-Install-Procedure.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
